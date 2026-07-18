"""Build Fishtown Web Design contract .docx with local embedded images."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "public" / "content" / "home"
SOURCE_MD = Path(r"c:\Users\Tre\Downloads\Fishtown_Web_Design_Contract.md")
OUTPUT = ROOT / "Fishtown_Web_Design_Contract.docx"

LOGO = ASSETS / "logo.png"
FOOTER_LOGO = ASSETS / "footer-logo.png"
FISH = ASSETS / "fish.png"


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    for level, size in ((1, 16), (2, 13)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_centered_picture(doc: Document, image_path: Path, width_inches: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_rich_text(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("["):
            label = token[1 : token.index("]")]
            url = token[token.index("(") + 1 : -1]
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0xC9, 0x7D, 0x60)
            run.underline = True
            paragraph.add_run(f" ({url})")
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold:
        run = paragraph.add_run(text)
        run.bold = True
    else:
        add_rich_text(paragraph, text)


def parse_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:\-|]+\|?\s*", line))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_rich_text(paragraph, cell_text)
            if r_idx == 0 or (c_idx == 0 and "**" in cell_text):
                for run in paragraph.runs:
                    run.bold = True


def extract_contract_body(md_text: str) -> str:
    start = md_text.find("# WEBSITE DESIGN, DEVELOPMENT, HOSTING & MAINTENANCE AGREEMENT")
    end = md_text.find("## SIGNATURES")
    if start == -1:
        raise ValueError("Could not find contract title in source markdown.")
    if end == -1:
        return md_text[start:].strip()
    return md_text[start:end].strip()


def convert_markdown_to_docx(md_text: str, doc: Document) -> None:
    lines = extract_contract_body(md_text).splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("<p align=") or line.strip().startswith("<img"):
            i += 1
            continue
        if line.strip() == "</p>":
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, line[2:].strip())
            i += 1
            continue

        if re.match(r"^\s{2}-\s", line):
            paragraph = doc.add_paragraph(style="List Bullet 2")
            add_rich_text(paragraph, line.strip()[2:].strip())
            i += 1
            continue

        add_paragraph(doc, line.strip())
        i += 1


def build_docx() -> Path:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")
    for image in (LOGO, FOOTER_LOGO, FISH):
        if not image.exists():
            raise FileNotFoundError(f"Missing asset: {image}")

    md_text = SOURCE_MD.read_text(encoding="utf-8")

    doc = Document()
    set_document_style(doc)

    add_centered_picture(doc, LOGO, 3.0)
    add_paragraph(
        doc,
        "Hand-coded websites for Philadelphia small businesses.",
        bold=True,
        center=True,
    )
    add_paragraph(doc, "Fast sites. Clean code. No shortcuts.", center=True)
    add_paragraph(
        doc,
        "fishtownwebdesign.com  ·  (717) 333-8691  ·  help@fishtownwebdesign.com",
        center=True,
    )
    doc.add_paragraph()

    convert_markdown_to_docx(md_text, doc)

    doc.add_page_break()
    add_centered_picture(doc, FOOTER_LOGO, 2.0)
    doc.add_paragraph()

    add_paragraph(doc, "DESIGNER: Fishtown Web Design LLC", bold=True)
    add_paragraph(doc, "Signed: ______________________________")
    add_paragraph(doc, "Printed Name: George Seibert")
    add_paragraph(doc, "Title: Founder")
    add_paragraph(doc, "Date: ______________________________")
    doc.add_paragraph()
    add_paragraph(doc, "CLIENT: [CLIENT LEGAL NAME]", bold=True)
    add_paragraph(doc, "Signed: ______________________________")
    add_paragraph(doc, "Printed Name: [NAME]")
    add_paragraph(doc, "Title: [TITLE]")
    add_paragraph(doc, "Date: ______________________________")
    doc.add_paragraph()

    add_centered_picture(doc, FISH, 0.5)
    add_paragraph(
        doc,
        "Fishtown Web Design LLC · Fishtown, Philadelphia, PA",
        center=True,
    )
    add_paragraph(
        doc,
        "www.fishtownwebdesign.com · (717) 333-8691 · help@fishtownwebdesign.com",
        center=True,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_docx()
    print(f"Created: {output}")
